import os
import time
import base64
import mimetypes
import re
import csv
import json
from datetime import datetime
from typing import Dict, List, Any, Optional, Tuple
import google.generativeai as genai
from PIL import Image
from docx import Document
import pandas as pd
from pathlib import Path

# === Config ===
genai.configure(api_key="AIzaSyA8ArcDpUOYYplu3gMK8Smdl2rYNUjQVmk")

# Load the Gemini multimodal model
model = genai.GenerativeModel(model_name="gemini-1.5-flash")

# Database table definitions with common fields
DATABASE_TABLES = {
    'patients_registration': ['patient_id', 'first_name', 'last_name', 'date_of_birth', 'gender', 'phone', 'email',
                              'address', 'emergency_contact'],
    'medical_history': ['patient_id', 'condition', 'diagnosis_date', 'status', 'notes', 'doctor_id'],
    'allergy_records': ['patient_id', 'allergen', 'reaction_type', 'severity', 'date_recorded'],
    'prescription': ['patient_id', 'medication', 'dosage', 'frequency', 'start_date', 'end_date', 'doctor_id'],
    'vitals_history': ['patient_id', 'blood_pressure_systolic', 'blood_pressure_diastolic', 'heart_rate', 'temperature',
                       'weight', 'height', 'date_recorded'],
    'bloodtests': ['patient_id', 'test_type', 'result_value', 'unit', 'reference_range', 'test_date', 'lab_id'],
    'diagnosis': ['patient_id', 'primary_diagnosis', 'secondary_diagnosis', 'icd_code', 'diagnosis_date', 'doctor_id'],
    'symptoms_checker': ['patient_id', 'symptom', 'severity', 'duration', 'date_reported'],
    'family_history': ['patient_id', 'relation', 'condition', 'age_of_onset', 'status'],
    'social_history': ['patient_id', 'smoking_status', 'alcohol_use', 'drug_use', 'exercise_frequency', 'occupation']
}


class MedicalDataProcessor:
    """Core medical data processing class - backend logic"""

    def __init__(self):
        self.processed_data = {}
        self.cleansing_patterns = self._setup_cleansing_patterns()

    def _setup_cleansing_patterns(self):
        """Setup regex patterns for data cleansing"""
        return {
            'phone': re.compile(r'(\d{3}[-.\s]?\d{3}[-.\s]?\d{4})'),
            'email': re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'),
            'date': re.compile(r'(\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\d{4}[/-]\d{1,2}[/-]\d{1,2})'),
            'blood_pressure': re.compile(r'(\d{2,3})/(\d{2,3})'),
            'temperature': re.compile(r'(\d{2,3}\.?\d?)°?[CF]?'),
            'weight': re.compile(r'(\d{2,3}\.?\d?)\s?(kg|lbs?|pounds?)'),
            'height': re.compile(r'(\d{1,2})\'\s?(\d{1,2})"?|(\d{3})\s?cm'),
            'medication': re.compile(r'\b[A-Z][a-z]+(?:cillin|pril|olol|statin|mycin)\b', re.IGNORECASE)
        }

    def cleanse_text(self, text: str) -> str:
        """Clean and standardize extracted text"""
        # Remove excessive whitespace and special characters
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'[^\w\s\-.,/():]', ' ', text)

        # Standardize common medical abbreviations
        abbreviations = {
            'w/': 'with',
            'w/o': 'without',
            'hx': 'history',
            'dx': 'diagnosis',
            'rx': 'prescription',
            'pt': 'patient',
            'dob': 'date of birth',
            'bp': 'blood pressure',
            'hr': 'heart rate',
            'temp': 'temperature'
        }

        for abbr, full in abbreviations.items():
            text = re.sub(rf'\b{abbr}\b', full, text, flags=re.IGNORECASE)

        return text.strip()

    def extract_structured_data(self, text: str) -> Dict[str, Any]:
        """Extract structured data using enhanced AI prompting"""

        structured_prompt = f"""
        Please analyze this medical document text and extract structured information. 
        Return the data in JSON format with the following categories:

        PATIENT_INFO: {{
            "name": "",
            "dob": "",
            "gender": "",
            "phone": "",
            "email": "",
            "address": "",
            "mrn": ""
        }},
        VITALS: {{
            "blood_pressure": "",
            "heart_rate": "",
            "temperature": "",
            "weight": "",
            "height": "",
            "date": ""
        }},
        MEDICATIONS: [{{
            "name": "",
            "dosage": "",
            "frequency": "",
            "start_date": "",
            "instructions": ""
        }}],
        ALLERGIES: [{{
            "allergen": "",
            "reaction": "",
            "severity": ""
        }}],
        DIAGNOSES: [{{
            "condition": "",
            "icd_code": "",
            "date": "",
            "status": ""
        }}],
        LAB_RESULTS: [{{
            "test_name": "",
            "result": "",
            "unit": "",
            "reference_range": "",
            "date": ""
        }}],
        SYMPTOMS: [{{
            "symptom": "",
            "severity": "",
            "duration": "",
            "date": ""
        }}],
        FAMILY_HISTORY: [{{
            "relation": "",
            "condition": "",
            "age_of_onset": ""
        }}],
        SOCIAL_HISTORY: {{
            "smoking": "",
            "alcohol": "",
            "occupation": "",
            "exercise": ""
        }}

        Medical Document Text:
        {text}

        Return only valid JSON without any additional text or formatting.
        """

        try:
            response = model.generate_content(structured_prompt)
            # Clean the response to extract JSON
            response_text = response.text.strip()

            # Remove code blocks if present
            if response_text.startswith('```'):
                response_text = response_text.split('```')[1]
                if response_text.startswith('json'):
                    response_text = response_text[4:]

            return json.loads(response_text)
        except Exception as e:
            print(f"Error in structured extraction: {e}")
            return self._fallback_extraction(text)

    def _fallback_extraction(self, text: str) -> Dict[str, Any]:
        """Fallback extraction using regex patterns"""
        data = {
            "PATIENT_INFO": {},
            "VITALS": {},
            "MEDICATIONS": [],
            "ALLERGIES": [],
            "DIAGNOSES": [],
            "LAB_RESULTS": [],
            "SYMPTOMS": [],
            "FAMILY_HISTORY": [],
            "SOCIAL_HISTORY": {}
        }

        # Extract basic patterns
        phone_match = self.cleansing_patterns['phone'].search(text)
        if phone_match:
            data["PATIENT_INFO"]["phone"] = phone_match.group(1)

        email_match = self.cleansing_patterns['email'].search(text)
        if email_match:
            data["PATIENT_INFO"]["email"] = email_match.group()

        bp_match = self.cleansing_patterns['blood_pressure'].search(text)
        if bp_match:
            data["VITALS"]["blood_pressure"] = f"{bp_match.group(1)}/{bp_match.group(2)}"

        return data

    def _create_comprehensive_notes(self, structured_data: Dict[str, Any], full_text: str) -> Dict[str, str]:
        """Create comprehensive notes from full text, including ALL document text"""

        # Clean and prepare the full document text
        full_document_text = self._clean_full_text(full_text)

        # Create base notes with FULL document content
        base_notes = {
            'patient_notes': full_document_text,
            'vitals_notes': full_document_text,
            'medication_notes': full_document_text,
            'allergy_notes': full_document_text,
            'diagnosis_notes': full_document_text,
            'lab_notes': full_document_text,
            'symptom_notes': full_document_text,
            'family_history_notes': full_document_text,
            'social_history_notes': full_document_text
        }

        # Add category-specific context at the beginning if found
        category_contexts = self._extract_category_contexts(full_text)

        # Prepend specific context to each category while keeping full text
        for category, context in category_contexts.items():
            if context and category in base_notes:
                base_notes[category] = f"[SPECIFIC CONTEXT]: {context}\n\n[FULL DOCUMENT]: {full_document_text}"

        return base_notes

    def _clean_full_text(self, text: str) -> str:
        """Clean the full text while preserving all important information"""
        # Remove excessive whitespace but keep paragraph structure
        cleaned_lines = []
        for line in text.split('\n'):
            line_clean = line.strip()
            if line_clean:
                cleaned_lines.append(line_clean)

        # Join with proper spacing
        cleaned_text = '\n'.join(cleaned_lines)

        # Remove any OCR artifacts but keep medical content
        cleaned_text = re.sub(r'\s+', ' ', cleaned_text)
        cleaned_text = re.sub(r'([.!?])\s*([A-Z])', r'\1 \2', cleaned_text)

        return cleaned_text.strip()

    def _extract_category_contexts(self, full_text: str) -> Dict[str, str]:
        """Extract category-specific context while preserving in full notes"""
        contexts = {}

        # Define category keywords and their contexts
        category_keywords = {
            'patient_notes': ['chief complaint', 'reason for visit', 'history of present illness', 'background',
                              'overview'],
            'vitals_notes': ['vital signs', 'physical examination', 'physical exam', 'assessment', 'measurements'],
            'medication_notes': ['medications', 'prescriptions', 'therapy', 'treatment plan', 'drug therapy'],
            'allergy_notes': ['allergies', 'allergic reactions', 'adverse reactions', 'sensitivities'],
            'diagnosis_notes': ['diagnosis', 'impression', 'findings', 'assessment', 'conclusion'],
            'lab_notes': ['laboratory results', 'lab results', 'test results', 'laboratory', 'pathology'],
            'symptom_notes': ['symptoms', 'complaints', 'presentation', 'manifestations'],
            'family_history_notes': ['family history', 'hereditary', 'genetic history', 'familial'],
            'social_history_notes': ['social history', 'lifestyle', 'habits', 'occupation', 'smoking', 'alcohol']
        }

        # Split text into sentences more robustly
        sentences = re.split(r'[.!?]+\s+', full_text.strip())
        sentences = [s.strip() for s in sentences if s.strip()]

        # Extract context for each category
        for category, keywords in category_keywords.items():
            context_sentences = []

            for sentence in sentences:
                sentence = sentence.strip()
                if any(keyword.lower() in sentence.lower() for keyword in keywords) and len(sentence) > 10:
                    try:
                        sentence_index = sentences.index(sentence)
                    except ValueError:
                        sentence_index = -1
                        sentence_lower = sentence.lower().strip()

                        for i, sent in enumerate(sentences):
                            if sentence_lower == sent.lower().strip():
                                sentence_index = i
                                break

                        if sentence_index == -1:
                            for i, sent in enumerate(sentences):
                                if (sentence_lower in sent.lower() and len(sentence) > 20) or \
                                        (sent.lower().strip() in sentence_lower and len(sent) > 20):
                                    sentence_index = i
                                    break

                        if sentence_index == -1:
                            print(f"Warning: Could not find sentence match: '{sentence[:50]}...'")
                            continue

                    context_part = sentence
                    if sentence_index + 1 < len(sentences):
                        context_part += '. ' + sentences[sentence_index + 1].strip()
                    context_sentences.append(context_part)

            if context_sentences:
                contexts[category] = '. '.join(context_sentences[:2])

        return contexts

    def convert_to_database_format(self, structured_data: Dict[str, Any], source_file: str, full_text: str,
                                   patient_id: str) -> Dict[str, List[Dict]]:
        """Convert structured data to database format with custom patient ID"""
        database_records = {}
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        # Create comprehensive notes from full text
        comprehensive_notes = self._create_comprehensive_notes(structured_data, full_text)

        # Patient Registration
        if structured_data.get("PATIENT_INFO"):
            patient_info = structured_data["PATIENT_INFO"]
            if any(patient_info.values()):
                name_parts = patient_info.get("name", "").split() if patient_info.get("name") else ["", ""]
                first_name = name_parts[0] if name_parts else ""
                last_name = " ".join(name_parts[1:]) if len(name_parts) > 1 else ""

                database_records["patients_registration"] = [{
                    "patient_id": patient_id,
                    "first_name": first_name,
                    "last_name": last_name,
                    "date_of_birth": patient_info.get("dob", ""),
                    "gender": patient_info.get("gender", ""),
                    "phone": patient_info.get("phone", ""),
                    "email": patient_info.get("email", ""),
                    "address": patient_info.get("address", ""),
                    "emergency_contact": "",
                    "medical_record_number": patient_info.get("mrn", ""),
                    "notes": comprehensive_notes.get("patient_notes", ""),
                    "source_file": source_file,
                    "processed_date": timestamp
                }]

        # Vitals History
        if structured_data.get("VITALS") and any(structured_data["VITALS"].values()):
            vitals = structured_data["VITALS"]
            bp_parts = vitals.get("blood_pressure", "").split("/") if vitals.get("blood_pressure") else ["", ""]

            database_records["vitals_history"] = [{
                "patient_id": patient_id,
                "blood_pressure_systolic": bp_parts[0] if len(bp_parts) > 0 else "",
                "blood_pressure_diastolic": bp_parts[1] if len(bp_parts) > 1 else "",
                "heart_rate": vitals.get("heart_rate", ""),
                "temperature": vitals.get("temperature", ""),
                "weight": vitals.get("weight", ""),
                "height": vitals.get("height", ""),
                "date_recorded": vitals.get("date", timestamp),
                "notes": comprehensive_notes.get("vitals_notes", ""),
                "source_file": source_file
            }]

        # Medications/Prescriptions
        if structured_data.get("MEDICATIONS"):
            database_records["prescription"] = []
            for med in structured_data["MEDICATIONS"]:
                if med.get("name"):
                    database_records["prescription"].append({
                        "patient_id": patient_id,
                        "medication": med.get("name", ""),
                        "dosage": med.get("dosage", ""),
                        "frequency": med.get("frequency", ""),
                        "start_date": med.get("start_date", timestamp),
                        "end_date": med.get("end_date", ""),
                        "doctor_id": "",
                        "instructions": med.get("instructions", ""),
                        "notes": comprehensive_notes.get("medication_notes", ""),
                        "source_file": source_file
                    })

        # Allergies
        if structured_data.get("ALLERGIES"):
            database_records["allergy_records"] = []
            for allergy in structured_data["ALLERGIES"]:
                if allergy.get("allergen"):
                    database_records["allergy_records"].append({
                        "patient_id": patient_id,
                        "allergen": allergy.get("allergen", ""),
                        "reaction_type": allergy.get("reaction", ""),
                        "severity": allergy.get("severity", ""),
                        "date_recorded": timestamp,
                        "notes": comprehensive_notes.get("allergy_notes", ""),
                        "source_file": source_file
                    })

        # Diagnoses
        if structured_data.get("DIAGNOSES"):
            database_records["diagnosis"] = []
            for diag in structured_data["DIAGNOSES"]:
                if diag.get("condition"):
                    database_records["diagnosis"].append({
                        "patient_id": patient_id,
                        "primary_diagnosis": diag.get("condition", ""),
                        "secondary_diagnosis": "",
                        "icd_code": diag.get("icd_code", ""),
                        "diagnosis_date": diag.get("date", timestamp),
                        "doctor_id": "",
                        "status": diag.get("status", "active"),
                        "notes": comprehensive_notes.get("diagnosis_notes", ""),
                        "source_file": source_file
                    })

        # Lab Results
        if structured_data.get("LAB_RESULTS"):
            database_records["bloodtests"] = []
            for lab in structured_data["LAB_RESULTS"]:
                if lab.get("test_name"):
                    database_records["bloodtests"].append({
                        "patient_id": patient_id,
                        "test_type": lab.get("test_name", ""),
                        "result_value": lab.get("result", ""),
                        "unit": lab.get("unit", ""),
                        "reference_range": lab.get("reference_range", ""),
                        "test_date": lab.get("date", timestamp),
                        "lab_id": "",
                        "notes": comprehensive_notes.get("lab_notes", ""),
                        "source_file": source_file
                    })

        # Symptoms
        if structured_data.get("SYMPTOMS"):
            database_records["symptoms_checker"] = []
            for symptom in structured_data["SYMPTOMS"]:
                if symptom.get("symptom"):
                    database_records["symptoms_checker"].append({
                        "patient_id": patient_id,
                        "symptom": symptom.get("symptom", ""),
                        "severity": symptom.get("severity", ""),
                        "duration": symptom.get("duration", ""),
                        "date_reported": symptom.get("date", timestamp),
                        "notes": comprehensive_notes.get("symptom_notes", ""),
                        "source_file": source_file
                    })

        # Family History
        if structured_data.get("FAMILY_HISTORY"):
            database_records["family_history"] = []
            for fh in structured_data["FAMILY_HISTORY"]:
                if fh.get("condition"):
                    database_records["family_history"].append({
                        "patient_id": patient_id,
                        "relation": fh.get("relation", ""),
                        "condition": fh.get("condition", ""),
                        "age_of_onset": fh.get("age_of_onset", ""),
                        "status": "reported",
                        "notes": comprehensive_notes.get("family_history_notes", ""),
                        "source_file": source_file
                    })

        # Social History
        if structured_data.get("SOCIAL_HISTORY") and any(structured_data["SOCIAL_HISTORY"].values()):
            social = structured_data["SOCIAL_HISTORY"]
            database_records["social_history"] = [{
                "patient_id": patient_id,
                "smoking_status": social.get("smoking", ""),
                "alcohol_use": social.get("alcohol", ""),
                "drug_use": "",
                "exercise_frequency": social.get("exercise", ""),
                "occupation": social.get("occupation", ""),
                "notes": comprehensive_notes.get("social_history_notes", ""),
                "source_file": source_file
            }]

        return database_records

    def save_to_csv(self, database_records: Dict[str, List[Dict]], csv_output_folder: str):
        """Save database records to consolidated CSV files by table type"""
        for table_name, records in database_records.items():
            if records:
                csv_filename = f"{table_name}.csv"
                csv_path = os.path.join(csv_output_folder, csv_filename)

                df = pd.DataFrame(records)

                if os.path.exists(csv_path):
                    existing_df = pd.read_csv(csv_path)
                    combined_df = pd.concat([existing_df, df], ignore_index=True)
                    combined_df.to_csv(csv_path, index=False)
                    print(f"✅ Appended {len(records)} records to {csv_path}")
                else:
                    df.to_csv(csv_path, index=False)
                    print(f"✅ Created {csv_path} with {len(records)} records")


class MedicalOCRInterface:
    """Frontend interface class for medical OCR processing"""

    def __init__(self):
        self.processor = MedicalDataProcessor()
        self.patient_id = None
        self.input_folder = None
        self.output_folder = None
        self.selected_files = []

    def validate_patient_id(self, patient_id: str) -> Tuple[bool, str]:
        """Validate patient ID input"""
        if not patient_id or not patient_id.strip():
            return False, "Patient ID cannot be empty"

        # Clean the patient ID
        cleaned_id = patient_id.strip()

        # Check if it contains only valid characters (alphanumeric, hyphens, underscores)
        if not re.match(r'^[a-zA-Z0-9_-]+$', cleaned_id):
            return False, "Patient ID can only contain letters, numbers, hyphens, and underscores"

        if len(cleaned_id) > 50:
            return False, "Patient ID must be 50 characters or less"

        return True, cleaned_id

    def validate_folder_path(self, folder_path: str) -> Tuple[bool, str]:
        """Validate folder path"""
        if not folder_path or not folder_path.strip():
            return False, "Folder path cannot be empty"

        path = Path(folder_path.strip())

        if not path.exists():
            return False, "Folder does not exist"

        if not path.is_dir():
            return False, "Path is not a directory"

        return True, str(path)

    def get_supported_image_files(self, folder_path: str) -> List[str]:
        """Get list of supported image files in folder"""
        supported_extensions = {'.png', '.jpg', '.jpeg', '.tiff', '.tif', '.bmp', '.webp'}
        image_files = []

        try:
            folder = Path(folder_path)
            for file in folder.iterdir():
                if file.is_file() and file.suffix.lower() in supported_extensions:
                    image_files.append(str(file))
        except Exception as e:
            print(f"Error reading folder: {e}")

        return sorted(image_files)

    def set_patient_id(self, patient_id: str) -> bool:
        """Set and validate patient ID"""
        is_valid, result = self.validate_patient_id(patient_id)
        if is_valid:
            self.patient_id = result
            return True
        else:
            print(f"Invalid Patient ID: {result}")
            return False

    def set_input_folder(self, folder_path: str) -> bool:
        """Set and validate input folder"""
        is_valid, result = self.validate_folder_path(folder_path)
        if is_valid:
            self.input_folder = result
            return True
        else:
            print(f"Invalid input folder: {result}")
            return False

    def set_output_folder(self, folder_path: str) -> bool:
        """Set and validate output folder"""
        is_valid, result = self.validate_folder_path(folder_path)
        if is_valid:
            self.output_folder = result
            # Create csv subfolder
            csv_folder = os.path.join(result, "csv_database_ready")
            os.makedirs(csv_folder, exist_ok=True)
            return True
        else:
            print(f"Invalid output folder: {result}")
            return False

    def get_available_files(self) -> List[Dict[str, str]]:
        """Get available image files with metadata"""
        if not self.input_folder:
            return []

        files = self.get_supported_image_files(self.input_folder)
        file_info = []

        for file_path in files:
            path = Path(file_path)
            try:
                stat = path.stat()
                size_mb = round(stat.st_size / (1024 * 1024), 2)
                modified = datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M")

                file_info.append({
                    'path': file_path,
                    'name': path.name,
                    'size_mb': size_mb,
                    'modified': modified
                })
            except Exception as e:
                print(f"Error getting file info for {file_path}: {e}")

        return file_info

    def set_selected_files(self, file_paths: List[str]) -> bool:
        """Set selected files for processing"""
        if not file_paths:
            print("No files selected")
            return False

        # Validate all files exist
        valid_files = []
        for file_path in file_paths:
            if Path(file_path).exists():
                valid_files.append(file_path)
            else:
                print(f"Warning: File not found: {file_path}")

        if not valid_files:
            print("No valid files selected")
            return False

        self.selected_files = valid_files
        return True

    def validate_processing_requirements(self) -> Tuple[bool, List[str]]:
        """Validate all requirements before processing"""
        errors = []

        if not self.patient_id:
            errors.append("Patient ID is required")
            
        if not self.output_folder:
            errors.append("Output folder is required")

        if not self.selected_files:
            errors.append("At least one file must be selected")

        return len(errors) == 0, errors

    def extract_text_from_image(self, image_path: str) -> Optional[str]:
        """Enhanced text extraction with better prompting"""
        mime_type, _ = mimetypes.guess_type(image_path)
        if not mime_type:
            mime_type = "image/png"

        try:
            with open(image_path, "rb") as f:
                image_bytes = f.read()

            response = model.generate_content([
                {"mime_type": mime_type, "data": image_bytes},
                {"text": """
                Extract ALL readable text from this medical document image with high accuracy. 
                Preserve the structure and formatting as much as possible.
                Pay special attention to:
                - Patient names, dates of birth, contact information
                - Medical record numbers, appointment dates
                - Vital signs (blood pressure, heart rate, temperature, weight, height)
                - Medications, dosages, and instructions
                - Diagnoses, ICD codes, and medical conditions
                - Lab results, test values, and reference ranges
                - Allergies and reactions
                - Symptoms and their descriptions
                - Family history information
                - Social history (smoking, alcohol, occupation)

                Format the output as clear, readable text maintaining the original document structure.
                """}
            ])
            return response.text.strip()
        except Exception as e:
            print(f"❌ Error processing {image_path}: {e}")
            return None

    def save_text_to_word(self, text: str, output_path: str):
        """Save extracted text to Word document"""
        try:
            doc = Document()
            doc.add_heading('Extracted Medical Document Text', 0)
            doc.add_paragraph(f'Patient ID: {self.patient_id}')
            doc.add_paragraph(f'Processing Date: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph('')  # Empty line

            for line in text.split('\n'):
                if line.strip():
                    doc.add_paragraph(line.strip())

            doc.save(output_path)
        except Exception as e:
            print(f"Error saving Word document: {e}")

    def process_files(self, progress_callback=None) -> Dict[str, Any]:
        """Process selected files and return results"""
        # Validate requirements
        is_valid, errors = self.validate_processing_requirements()
        if not is_valid:
            return {
                'success': False,
                'error': f"Processing requirements not met: {'; '.join(errors)}",
                'results': {}
            }

        results = {
            'success': True,
            'total_files': len(self.selected_files),
            'processed_files': 0,
            'total_records': 0,
            'files_processed': [],
            'files_failed': [],
            'csv_files_created': []
        }

        csv_output_folder = os.path.join(self.output_folder, "csv_database_ready")

        print("🏥 Starting Enhanced Medical OCR Processing...")
        print("=" * 60)
        print(f"📋 Patient ID: {self.patient_id}")
        print(f"📁 Input Folder: {self.input_folder}")
        print(f"📁 Output Folder: {self.output_folder}")
        print(f"📄 Files to Process: {len(self.selected_files)}")

        for i, file_path in enumerate(self.selected_files):
            filename = Path(file_path).name

            # Call progress callback if provided
            if progress_callback:
                progress = (i / len(self.selected_files)) * 100
                progress_callback(progress, f"Processing {filename}")

            print(f"\n🔍 Processing: {filename}")
            print("-" * 40)

            try:
                # Extract text
                extracted_text = self.extract_text_from_image(file_path)

                if extracted_text:
                    # Save original extraction to Word
                    base_filename = Path(file_path).stem
                    word_output_path = os.path.join(self.output_folder, f"{base_filename}_extracted.docx")
                    self.save_text_to_word(extracted_text, word_output_path)
                    print(f"📄 Word document saved: {word_output_path}")

                    # Cleanse the text
                    cleansed_text = self.processor.cleanse_text(extracted_text)

                    # Extract structured data
                    print("🧠 Extracting structured data...")
                    structured_data = self.processor.extract_structured_data(cleansed_text)

                    # Convert to database format with custom patient ID
                    print("🗄️  Converting to database format...")
                    database_records = self.processor.convert_to_database_format(
                        structured_data, filename, cleansed_text, self.patient_id
                    )

                    # Save to CSV
                    if database_records:
                        self.processor.save_to_csv(database_records, csv_output_folder)
                        record_count = sum(len(records) for records in database_records.values())
                        results['total_records'] += record_count
                        print(f"📊 Generated {record_count} total database records")

                        # Track CSV files created
                        for table_name in database_records.keys():
                            csv_file = f"{table_name}.csv"
                            if csv_file not in results['csv_files_created']:
                                results['csv_files_created'].append(csv_file)
                    else:
                        print("⚠️  No structured data found to convert")

                    results['processed_files'] += 1
                    results['files_processed'].append({
                        'filename': filename,
                        'records_created': record_count if database_records else 0,
                        'word_file': word_output_path
                    })

                else:
                    print(f"❌ No text extracted from: {filename}")
                    results['files_failed'].append({
                        'filename': filename,
                        'error': 'No text extracted'
                    })

            except Exception as e:
                print(f"❌ Error processing {filename}: {e}")
                results['files_failed'].append({
                    'filename': filename,
                    'error': str(e)
                })

        # Final progress callback
        if progress_callback:
            progress_callback(100, "Processing complete")

        print("\n" + "=" * 60)
        print("🎉 PROCESSING COMPLETE!")
        print(f"📁 Files processed: {results['processed_files']}")
        print(f"📊 Total database records created: {results['total_records']}")
        print(f"💾 CSV files saved to: {csv_output_folder}")
        print(f"📄 Word documents saved to: {self.output_folder}")
        print("\n📋 CSV files are ready for database import!")
        print("=" * 60)

        return results


# Frontend Helper Functions
def get_user_input_patient_id() -> Optional[str]:
    """Console-based patient ID input (for testing)"""
    while True:
        patient_id = input("\n📋 Enter Patient ID (alphanumeric, hyphens, underscores allowed): ").strip()
        if not patient_id:
            print("❌ Patient ID cannot be empty. Please try again.")
            continue

        interface = MedicalOCRInterface()
        if interface.set_patient_id(patient_id):
            return patient_id
        else:
            print("❌ Invalid Patient ID. Please try again.")


def get_user_input_folder(folder_type: str) -> Optional[str]:
    """Console-based folder input (for testing)"""
    while True:
        folder_path = input(f"\n📁 Enter {folder_type} folder path: ").strip()
        if not folder_path:
            print(f"❌ {folder_type} folder path cannot be empty. Please try again.")
            continue

        interface = MedicalOCRInterface()
        is_valid, result = interface.validate_folder_path(folder_path)
        if is_valid:
            return result
        else:
            print(f"❌ {result}. Please try again.")


def get_user_file_selection(input_folder: str) -> List[str]:
    """Console-based file selection (for testing)"""
    interface = MedicalOCRInterface()
    interface.set_input_folder(input_folder)

    available_files = interface.get_available_files()

    if not available_files:
        print("❌ No supported image files found in the folder.")
        return []

    print(f"\n📄 Found {len(available_files)} supported image files:")
    print("-" * 60)

    for i, file_info in enumerate(available_files, 1):
        print(f"{i:2d}. {file_info['name']} ({file_info['size_mb']} MB) - {file_info['modified']}")

    while True:
        selection = input(
            f"\n📋 Enter file numbers to process (1-{len(available_files)}, comma-separated, or 'all' for all files): ").strip()

        if selection.lower() == 'all':
            return [f['path'] for f in available_files]

        try:
            indices = [int(x.strip()) - 1 for x in selection.split(',')]
            selected_files = []

            for idx in indices:
                if 0 <= idx < len(available_files):
                    selected_files.append(available_files[idx]['path'])
                else:
                    print(f"❌ Invalid file number: {idx + 1}")
                    break
            else:
                if selected_files:
                    return selected_files

        except ValueError:
            print("❌ Invalid input. Please enter numbers separated by commas or 'all'.")


def console_progress_callback(progress: float, message: str):
    """Simple console progress callback"""
    print(f"📈 Progress: {progress:.1f}% - {message}")


# Example usage for console-based testing
def main_console_interface():
    """Main console interface for testing"""
    print("🏥 Medical Document OCR Processor")
    print("=" * 50)

    # Initialize interface
    interface = MedicalOCRInterface()

    # Get patient ID
    patient_id = get_user_input_patient_id()
    if not patient_id:
        print("❌ Failed to get valid patient ID. Exiting.")
        return

    interface.set_patient_id(patient_id)

    # Get input folder
    input_folder = get_user_input_folder("input")
    if not input_folder:
        print("❌ Failed to get valid input folder. Exiting.")
        return

    interface.set_input_folder(input_folder)

    # Get output folder
    output_folder = get_user_input_folder("output")
    if not output_folder:
        print("❌ Failed to get valid output folder. Exiting.")
        return

    interface.set_output_folder(output_folder)

    # Get file selection
    selected_files = get_user_file_selection(input_folder)
    if not selected_files:
        print("❌ No files selected. Exiting.")
        return

    interface.set_selected_files(selected_files)

    # Confirm processing
    print(f"\n📋 Processing Summary:")
    print(f"   Patient ID: {patient_id}")
    print(f"   Input Folder: {input_folder}")
    print(f"   Output Folder: {output_folder}")
    print(f"   Files to Process: {len(selected_files)}")

    confirm = input("\n🚀 Start processing? (y/n): ").strip().lower()
    if confirm != 'y':
        print("❌ Processing cancelled.")
        return

    # Process files
    results = interface.process_files(progress_callback=console_progress_callback)

    # Display results
    if results['success']:
        print(f"\n✅ Processing completed successfully!")
        print(f"   Files processed: {results['processed_files']}/{results['total_files']}")
        print(f"   Total records: {results['total_records']}")
        print(f"   CSV files created: {', '.join(results['csv_files_created'])}")

        if results['files_failed']:
            print(f"\n⚠️  Failed files:")
            for failed_file in results['files_failed']:
                print(f"   - {failed_file['filename']}: {failed_file['error']}")
    else:
        print(f"\n❌ Processing failed: {results.get('error', 'Unknown error')}")


# Example usage for GUI integration
class GUIIntegrationExample:
    """Example showing how to integrate with a GUI framework"""

    def __init__(self):
        self.interface = MedicalOCRInterface()

    def on_patient_id_changed(self, patient_id: str) -> bool:
        """Called when user enters patient ID in GUI"""
        return self.interface.set_patient_id(patient_id)

    def on_input_folder_selected(self, folder_path: str) -> bool:
        """Called when user selects input folder in GUI"""
        success = self.interface.set_input_folder(folder_path)
        if success:
            # Update file list in GUI
            files = self.interface.get_available_files()
            # Update GUI with available files
            self.update_file_list_in_gui(files)
        return success

    def on_output_folder_selected(self, folder_path: str) -> bool:
        """Called when user selects output folder in GUI"""
        return self.interface.set_output_folder(folder_path)

    def on_files_selected(self, file_paths: List[str]) -> bool:
        """Called when user selects files in GUI"""
        return self.interface.set_selected_files(file_paths)

    def on_process_button_clicked(self):
        """Called when user clicks process button in GUI"""
        # Validate requirements
        is_valid, errors = self.interface.validate_processing_requirements()
        if not is_valid:
            # Show error in GUI
            self.show_error_in_gui(f"Cannot start processing: {'; '.join(errors)}")
            return

        # Start processing (should be in separate thread for GUI)
        results = self.interface.process_files(progress_callback=self.gui_progress_callback)

        # Show results in GUI
        self.show_results_in_gui(results)

    def gui_progress_callback(self, progress: float, message: str):
        """Update progress bar and status in GUI"""
        # Update GUI progress bar and status label
        pass

    def update_file_list_in_gui(self, files: List[Dict[str, str]]):
        """Update file list in GUI"""
        # Populate file list widget with available files
        pass

    def show_error_in_gui(self, error_message: str):
        """Show error message in GUI"""
        # Display error dialog or status message
        pass

    def show_results_in_gui(self, results: Dict[str, Any]):
        """Show processing results in GUI"""
        # Display results dialog or update status
        pass


if __name__ == "__main__":
    # Run console interface for testing
    main_console_interface()